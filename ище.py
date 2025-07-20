from docx import Document

TEMPLATE_PATH = "template.docx"
OUTPUT_PATH = "output.docx"

context = {
    "[дата]": "20.07.2025",
    "[номер_ЕРДР]": "123456789",
    "[дата_регистрации ЕРДР]": "18.07.2025",
    "[статья_УК_РК]": "188 ч.1",
    "[фабула_уголовного_дела]": "н/л совершил тайное хищение товара в магазине Magnum",
    "[ФИО_потерпевшего]": "Омарова Балжан Сериковна",
    "[иниц_потерпевшего]": "Омарова Б.С.",
    "[г.р._потерпевшего]": "1985",
    "[м.ж._потерпевшего]": "г.Астана, пр.Туран 50/2",
    "[сумма_ущерба]": "8950",
    "[ФИО_подозреваемого]": "Иванов Николай Петрович",
    "[иниц_подозреваемого]": "Иванов Н.П.",
    "[г.р._подозреваемого]": "1990",
    "[м.ж._подозреваемого]": "г.Астана, ул. Сарыарка 10"
}

def replace_runs(paragraph, data):
    for run in paragraph.runs:
        for key, value in data.items():
            if key in run.text:
                run.text = run.text.replace(key, value)

def replace_all(doc, data):
    for paragraph in doc.paragraphs:
        replace_runs(paragraph, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_runs(paragraph, data)

def main():
    doc = Document(TEMPLATE_PATH)
    replace_all(doc, context)
    doc.save(OUTPUT_PATH)
    print(f"✅ Документ сохранён: {OUTPUT_PATH}")

if __name__ == "__main__":
    main()
